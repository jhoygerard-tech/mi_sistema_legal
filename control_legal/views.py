import io
import json
from datetime import datetime
from xhtml2pdf import pisa
from django.http import JsonResponse
from decimal import Decimal
from . import liquidacion_asistencia_services as asistencia_services
from django.contrib import messages
from django.contrib.auth import authenticate, login
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from django.http import FileResponse, HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.template.loader import get_template
from django.utils import timezone 
from . import services as srv 
from django.views.decorators.http import require_POST
from .services import explicar_en_idioma_nativo

from .forms import (CasoForm, ClienteForm, ExpedienteForm, DocumentoForm, MovimientoForm, EventoForm, TareaForm, EvidenciaForm, PlantillaContratoForm, GenerarContratoForm, ComentarioTareaForm)
from .models import Cliente, Expediente, Documento, Movimiento, Evento, Tarea, EvidenciaTarea, PlantillaContrato, Contrato, PlantillaMemorial, Memorial, PrecedenteJurisprudencial, ComentarioTarea, NotificacionInterna 
@login_required
def dashboard(request):
    from django.utils import timezone
    from datetime import timedelta

    hoy      = timezone.now().date()
    semana   = hoy + timedelta(days=7)
    dos_dias = hoy + timedelta(days=2)
    total        = Expediente.objects.count()
    activos      = Expediente.objects.filter(activo=True).count()
    concluidos   = Expediente.objects.filter(estado='concluido').count()
    casos_recientes = Expediente.objects.filter(
        activo=True).select_related('cliente', 'abogado_asignado').order_by('-id')[:5]

    casos_urgentes  = Expediente.objects.filter(
        activo=True,
        eventos__fecha_hora__date__lte=dos_dias,
        eventos__fecha_hora__date__gte=hoy,
    ).distinct().count()

    casos_proximos  = Expediente.objects.filter(
        activo=True,
        eventos__fecha_hora__date__gt=dos_dias,
        eventos__fecha_hora__date__lte=semana,
    ).distinct().count()

    casos_al_dia    = activos - casos_urgentes - casos_proximos

    audiencias_semana = Evento.objects.filter(
        fecha_hora__date__gte=hoy,
        fecha_hora__date__lte=semana,
    ).count()

    eventos_hoy = Evento.objects.filter(
        fecha_hora__date=hoy
    ).select_related('expediente__cliente').order_by('fecha_hora')

    tareas_pendientes = Tarea.objects.filter(
        estado__in=['pendiente', 'en_proceso']
    ).count()

    tareas_vencidas = Tarea.objects.filter(
        estado__in=['pendiente', 'en_proceso'],
        fecha_limite__lt=hoy,
    ).count()

    total_clientes = Cliente.objects.filter(activo=True).count()

    return render(request, 'control_legal/dashboard.html', {
        'total':            total,
        'activos':          activos,
        'concluidos':       concluidos,
        'casos_recientes':  casos_recientes,
        'casos_al_dia':     max(casos_al_dia, 0),
        'casos_proximos':   casos_proximos,
        'casos_urgentes':   casos_urgentes,
        'audiencias_semana': audiencias_semana,
        'eventos_hoy':       eventos_hoy,
        'tareas_pendientes': tareas_pendientes,
        'tareas_vencidas':   tareas_vencidas,
        'total_clientes':    total_clientes,
    })

# ── Listar expedientes ────────────────────────────────────────
@login_required
def listar_casos(request):
    busqueda       = request.GET.get('buscar', '')
    criterio       = request.GET.get('orden', 'id')
    estado_filtro  = request.GET.get('estado', '')
    materia_filtro = request.GET.get('materia', '')

    casos = Expediente.objects.filter(activo=True).select_related(
        'cliente', 'abogado_asignado'
    )

    # ── Búsqueda de texto ────────────────────────────────────────
    if busqueda:
        casos = casos.filter(
            Q(cliente__nombre_completo__icontains=busqueda) |
            Q(nurej__icontains=busqueda)                    |
            Q(materia__icontains=busqueda)                  |
            Q(tipo_proceso__icontains=busqueda)             |
            # CORRECCIÓN: buscar por nombre completo Y por username
            Q(abogado_asignado__first_name__icontains=busqueda) |
            Q(abogado_asignado__last_name__icontains=busqueda)  |
            Q(abogado_asignado__username__icontains=busqueda)
        )

    # ── Filtro por estado ────────────────────────────────────────
    if estado_filtro:
        casos = casos.filter(estado=estado_filtro)

    # ── Filtro por materia ───────────────────────────────────────
    if materia_filtro:
        casos = casos.filter(materia=materia_filtro)

    # ── Ordenamiento ─────────────────────────────────────────────
    orden = {
        'nombre':  'cliente__nombre_completo',
        'fecha':   '-fecha_inicio',
        'materia': 'materia',
        'id':      'id',
    }
    casos = casos.order_by(orden.get(criterio, '-fecha_inicio'))

    return render(request, 'control_legal/listar_casos.html', {
        'casos':          casos,
        'busqueda':       busqueda,
        'estado_filtro':  estado_filtro,
        'materia_filtro': materia_filtro,
    })

# ── Listar clientes ───────────────────────────────────────────
@login_required
def listar_clientes(request):
    busqueda = request.GET.get('buscar', '')
    clientes = Cliente.objects.filter(activo=True)
    if busqueda:
        clientes = clientes.filter(
            Q(nombre_completo__icontains=busqueda) |
            Q(ci__icontains=busqueda)              |
            Q(telefono__icontains=busqueda)
        )
    return render(request, 'control_legal/listar_clientes.html', {
        'clientes': clientes.order_by('nombre_completo'),
        'busqueda': busqueda,
    })


# ── Registrar cliente ─────────────────────────────────────────
@login_required
def registrar_cliente(request):
    form = ClienteForm(request.POST or None)
    if form.is_valid():
        form.save()
        messages.success(request, 'Cliente registrado exitosamente.')
        return redirect('listar_clientes')
    return render(request, 'control_legal/registrar_cliente.html', {'form': form})


# ── Editar cliente ────────────────────────────────────────────
@login_required
def editar_cliente(request, pk):
    cliente = get_object_or_404(Cliente, pk=pk)
    form = ClienteForm(request.POST or None, instance=cliente)
    if form.is_valid():
        form.save()
        messages.success(request, 'Cliente actualizado exitosamente.')
        return redirect('detalle_cliente', pk=pk)
    return render(request, 'control_legal/registrar_cliente.html', {
        'form': form, 'editando': True, 'cliente': cliente,
    })


# ── Detalle del cliente (con sus expedientes) ─────────────────
@login_required
def detalle_cliente(request, pk):
    cliente     = get_object_or_404(Cliente, pk=pk)
    expedientes = cliente.expedientes.filter(activo=True).order_by('-id')
    return render(request, 'control_legal/detalle_cliente.html', {
        'cliente':     cliente,
        'expedientes': expedientes,
    })


# ── Registrar expediente (desde el cliente) ───────────────────
@login_required
def registrar_caso(request):
    cliente_id = request.GET.get('cliente')
    cliente    = get_object_or_404(Cliente, pk=cliente_id) if cliente_id else None
    form       = ExpedienteForm(request.POST or None)
    if form.is_valid():
        expediente         = form.save(commit=False)
        expediente.cliente = cliente
        expediente.save()
        messages.success(request, 'Expediente registrado exitosamente.')
        if cliente:
            return redirect('detalle_cliente', pk=cliente.pk)
        return redirect('listar_casos')
    return render(request, 'control_legal/registrar_caso.html', {
        'form': form, 'cliente': cliente,
    })


# ── Editar expediente ─────────────────────────────────────────
@login_required
def editar_caso(request, pk):
    caso = get_object_or_404(Expediente, pk=pk)
    form = ExpedienteForm(request.POST or None, instance=caso)
    if form.is_valid():
        form.save()
        messages.success(request, 'Expediente actualizado exitosamente.')
        return redirect('detalle_caso', pk=pk)
    return render(request, 'control_legal/registrar_caso.html', {
        'form': form, 'editando': True,
    })


# ── Detalle del expediente (documentos + movimientos) ─────────
@login_required
def detalle_caso(request, pk):
    expediente = get_object_or_404(Expediente, pk=pk)
    doc_form   = DocumentoForm(request.POST or None, request.FILES or None)
    mov_form   = MovimientoForm(request.POST or None)

    if 'subir_documento' in request.POST and doc_form.is_valid():
        doc          = doc_form.save(commit=False)
        doc.expediente  = expediente
        doc.subido_por  = request.user
        doc.save()
        messages.success(request, 'Documento subido correctamente.')
        return redirect('detalle_caso', pk=pk)

    if 'agregar_movimiento' in request.POST and mov_form.is_valid():
        mov               = mov_form.save(commit=False)
        mov.expediente    = expediente
        mov.registrado_por = request.user
        mov.save()
        messages.success(request, 'Movimiento registrado.')
        return redirect('detalle_caso', pk=pk)

    tareas = expediente.tareas.select_related('asignada_a').order_by('fecha_limite')

    # Construir contexto del expediente para el Asistente IA
    # Este resumen se inyecta en el system prompt de Gemini
    movimientos_texto = '\n'.join([
        f"- {m.fecha.strftime('%d/%m/%Y')} | {m.get_tipo_display()}: {m.descripcion}"
        for m in expediente.movimientos.all()[:10]
    ]) or 'Sin movimientos registrados.'

    documentos_texto = '\n'.join([
        f"- {d.titulo} ({d.get_tipo_display()}) — {d.fecha_subida.strftime('%d/%m/%Y')}"
        for d in expediente.documentos.all()[:10]
    ]) or 'Sin documentos adjuntos.'

    tareas_texto = '\n'.join([
        f"- {'⚠️ VENCIDA' if t.esta_vencida else 'Vence'} {t.fecha_limite.strftime('%d/%m/%Y')}: {t.titulo} [{t.get_estado_display()}]"
        for t in tareas
    ]) or 'Sin tareas asignadas.'

    contexto_expediente = f"""EXPEDIENTE ACTIVO EN EL SISTEMA
================================
N° Expediente : {expediente.numero_expediente or '—'}
NUREJ / IANUS : {expediente.nurej or '—'}
Juzgado       : {expediente.juzgado or '—'}
Materia       : {expediente.materia}
Tipo de proceso: {expediente.tipo_proceso or '—'}
Estado        : {expediente.get_estado_display()}
Fecha inicio  : {expediente.fecha_inicio.strftime('%d/%m/%Y')}
Cliente       : {expediente.cliente.nombre_completo}
Abogado       : {str(expediente.abogado_asignado) if expediente.abogado_asignado else '—'}
Descripción   : {expediente.descripcion or '—'}

MOVIMIENTOS PROCESALES (últimos 10):
{movimientos_texto}

DOCUMENTOS ADJUNTOS:
{documentos_texto}

TAREAS PENDIENTES:
{tareas_texto}
================================
Usa este contexto para responder con precisión. Cuando el abogado
pida redactar un escrito, usa los datos reales del expediente.
Cita siempre el artículo legal correspondiente."""

    return render(request, 'control_legal/detalle_caso.html', {
        'expediente':              expediente,
        'documentos':              expediente.documentos.all(),
        'movimientos':             expediente.movimientos.all(),
        'doc_form':                doc_form,
        'mov_form':                mov_form,
        'tareas_del_caso':         tareas,
        'mostrar_calculadora_familia': expediente.materia == 'Familia',
        'contexto_expediente':     contexto_expediente,
    })


# ── Eliminar expediente (lógico) ──────────────────────────────
@login_required
def eliminar_caso(request, pk):
    caso        = get_object_or_404(Expediente, pk=pk)
    caso.activo = False
    caso.save()
    messages.success(request, 'Expediente movido a la papelera.')
    return redirect('listar_casos')


# ── Papelera ──────────────────────────────────────────────────
@login_required
def papelera_casos(request):
    casos = Expediente.objects.filter(activo=False).order_by('-id')
    return render(request, 'control_legal/papelera.html', {'casos': casos})


# ── Restaurar desde papelera ──────────────────────────────────
@login_required
def restaurar_caso(request, pk):
    caso        = get_object_or_404(Expediente, pk=pk)
    caso.activo = True
    caso.save()
    messages.success(request, 'Expediente restaurado exitosamente.')
    return redirect('papelera_casos')


# ── Reporte PDF general ───────────────────────────────────────
@login_required
def exportar_pdf(request):
    """
    Informe ejecutivo completo del bufete:
    - Portada con fecha y generador
    - KPIs principales (expedientes, clientes, tareas, audiencias)
    - Desglose por materia, estado y abogado
    - Tabla detallada de expedientes activos
    - Resumen de tareas pendientes y vencidas
    - Próximas audiencias
    - Contratos y memoriales generados
    - Pie con firma del generador
    """
    from django.db.models import Count, Q
    from django.utils import timezone
    from datetime import timedelta

    hoy     = timezone.now().date()
    mes_act = hoy.month
    anio    = hoy.year
    hace_30 = hoy - timedelta(days=30)

    # ── Expedientes ──────────────────────────────────────────
    expedientes_activos    = Expediente.objects.filter(activo=True)
    total_expedientes      = Expediente.objects.count()
    total_activos          = expedientes_activos.filter(estado='activo').count()
    total_concluidos       = Expediente.objects.filter(estado='concluido').count()
    total_archivados       = Expediente.objects.filter(estado='archivado').count()
    total_abandono         = Expediente.objects.filter(estado='abandono').count()

    nuevos_mes = Expediente.objects.filter(
        fecha_inicio__month=mes_act, fecha_inicio__year=anio
    ).count()

    # Desglose por materia
    por_materia = (
        Expediente.objects
        .filter(activo=True)
        .values('materia')
        .annotate(total=Count('id'))
        .order_by('-total')
    )

    # Desglose por abogado
    por_abogado = (
        Expediente.objects
        .filter(activo=True, abogado_asignado__isnull=False)
        .values(
            'abogado_asignado__first_name',
            'abogado_asignado__last_name',
            'abogado_asignado__username',
        )
        .annotate(total=Count('id'))
        .order_by('-total')[:8]
    )
    por_abogado_lista = [
        {
            'nombre': (
                f"{r['abogado_asignado__first_name']} {r['abogado_asignado__last_name']}".strip()
                or r['abogado_asignado__username']
            ),
            'total': r['total'],
        }
        for r in por_abogado
    ]
    sin_abogado = expedientes_activos.filter(abogado_asignado__isnull=True).count()

    # ── Clientes ─────────────────────────────────────────────
    total_clientes  = Cliente.objects.filter(activo=True).count()
    nuevos_clientes = Cliente.objects.filter(
        fecha_registro__month=mes_act,
        fecha_registro__year=anio,
    ).count()

    # ── Tareas ───────────────────────────────────────────────
    tareas_pendientes  = Tarea.objects.filter(estado='pendiente').count()
    tareas_en_proceso  = Tarea.objects.filter(estado='en_proceso').count()
    tareas_completadas = Tarea.objects.filter(estado='completada').count()
    tareas_vencidas    = Tarea.objects.filter(
        estado__in=['pendiente', 'en_proceso'],
        fecha_limite__lt=hoy,
    ).count()

    tareas_detalle = (
        Tarea.objects
        .filter(estado__in=['pendiente', 'en_proceso'])
        .select_related('asignada_a', 'expediente__cliente')
        .order_by('fecha_limite')[:10]
    )

    # ── Agenda / Audiencias ───────────────────────────────────
    proximas_audiencias = (
        Evento.objects
        .filter(fecha_hora__date__gte=hoy)
        .select_related('expediente__cliente', 'responsable')
        .order_by('fecha_hora')[:8]
    )
    audiencias_mes = Evento.objects.filter(
        fecha_hora__year=anio,
        fecha_hora__month=mes_act,
    ).count()

    # ── Producción IA ─────────────────────────────────────────
    total_contratos  = Contrato.objects.count()
    contratos_mes    = Contrato.objects.filter(
        fecha_generado__month=mes_act,
        fecha_generado__year=anio,
    ).count()
    total_memoriales = Memorial.objects.count()
    memoriales_mes   = Memorial.objects.filter(
        fecha_generado__month=mes_act,
        fecha_generado__year=anio,
    ).count()

    # ── Tabla principal de expedientes ────────────────────────
    casos_tabla = (
        Expediente.objects
        .filter(activo=True)
        .select_related('cliente', 'abogado_asignado')
        .order_by('cliente__nombre_completo')
    )

    # ── Barras proporcionales para gráfico de materias ────────
    max_materia = max((r['total'] for r in por_materia), default=1)
    materias_grafico = [
        {
            'materia': r['materia'],
            'total':   r['total'],
            'pct':     int(r['total'] * 100 / max_materia),
        }
        for r in por_materia
    ]

    # ── Porcentajes de estados para gráfico pastel (SVG) ─────
    if total_expedientes:
        pct_activo    = round(total_activos    * 100 / total_expedientes)
        pct_concluido = round(total_concluidos * 100 / total_expedientes)
        pct_archivado = round(total_archivados * 100 / total_expedientes)
        pct_abandono  = round(total_abandono   * 100 / total_expedientes)
    else:
        pct_activo = pct_concluido = pct_archivado = pct_abandono = 0

    template = get_template('control_legal/reporte_pdf.html')
    html = template.render({
        'fecha':              datetime.now(),
        'generador':          request.user.get_full_name() or request.user.username,
        'hoy':                hoy,
        # Expedientes
        'total_expedientes':  total_expedientes,
        'total_activos':      total_activos,
        'total_concluidos':   total_concluidos,
        'total_archivados':   total_archivados,
        'total_abandono':     total_abandono,
        'nuevos_mes':         nuevos_mes,
        'sin_abogado':        sin_abogado,
        'casos_tabla':        casos_tabla,
        # Gráficos
        'materias_grafico':   materias_grafico,
        'por_abogado_lista':  por_abogado_lista,
        'pct_activo':         pct_activo,
        'pct_concluido':      pct_concluido,
        'pct_archivado':      pct_archivado,
        'pct_abandono':       pct_abandono,
        # Clientes
        'total_clientes':     total_clientes,
        'nuevos_clientes':    nuevos_clientes,
        # Tareas
        'tareas_pendientes':  tareas_pendientes,
        'tareas_en_proceso':  tareas_en_proceso,
        'tareas_completadas': tareas_completadas,
        'tareas_vencidas':    tareas_vencidas,
        'tareas_detalle':     tareas_detalle,
        # Agenda
        'proximas_audiencias': proximas_audiencias,
        'audiencias_mes':      audiencias_mes,
        # Producción IA
        'total_contratos':    total_contratos,
        'contratos_mes':      contratos_mes,
        'total_memoriales':   total_memoriales,
        'memoriales_mes':     memoriales_mes,
        'request':            request,
    })

    result = io.BytesIO()
    pisa.CreatePDF(io.BytesIO(html.encode('UTF-8')), dest=result, encoding='UTF-8')
    result.seek(0)
    nombre = f"Informe_Ejecutivo_LexNova_{datetime.now().strftime('%Y%m%d')}.pdf"
    return FileResponse(result, as_attachment=True, filename=nombre)

# ── Reporte PDF de un expediente ──────────────────────────────
@login_required
def reporte_un_caso(request, pk):
    caso     = get_object_or_404(Expediente, pk=pk)
    template = get_template('control_legal/reporte_individual.html')
    html     = template.render({'caso': caso, 'fecha': datetime.now(), 'request': request})
    result   = io.BytesIO()
    pisa.pisaDocument(io.BytesIO(html.encode('UTF-8')), result, encoding='UTF-8')
    result.seek(0)
    return FileResponse(result, as_attachment=False, filename=f'Expediente_{caso.nurej}.pdf')


# ── Login ─────────────────────────────────────────────────────
def login_view(request):
    if request.user.is_authenticated:
        return redirect('dashboard')
    if request.method == 'POST':
        user = authenticate(
            request,
            username=request.POST.get('username'),
            password=request.POST.get('password'),
        )
        if user is not None:
            login(request, user)
            return redirect('dashboard')
        messages.error(request, 'Usuario o contraseña incorrectos.')
    return render(request, 'login.html')

# ── Agenda: lista de eventos ──────────────────────────────────
@login_required
def agenda(request):
    from django.utils import timezone
    from datetime import timedelta
    hoy       = timezone.now().date()
    eventos   = Evento.objects.filter(fecha_hora__date__gte=hoy).order_by('fecha_hora')
    urgentes  = eventos.filter(fecha_hora__date__lte=hoy + timedelta(days=2))
    return render(request, 'control_legal/agenda.html', {
        'eventos':  eventos,
        'urgentes': urgentes,
        'hoy':      hoy,
    })


# ── Agenda: registrar evento ──────────────────────────────────
@login_required
def registrar_evento(request):
    expediente_id = request.GET.get('expediente')
    initial = {}
    if expediente_id:
        try:
            initial['expediente'] = Expediente.objects.get(pk=expediente_id)
        except Expediente.DoesNotExist:
            pass
    form = EventoForm(request.POST or None, initial=initial)
    if form.is_valid():
        evento             = form.save(commit=False)
        evento.responsable = request.user
        evento.save()
        messages.success(request, 'Evento agendado correctamente.')
        return redirect('agenda')
    return render(request, 'control_legal/registrar_evento.html', {'form': form})


# ── Agenda: editar evento ─────────────────────────────────────
@login_required
def editar_evento(request, pk):
    evento = get_object_or_404(Evento, pk=pk)
    form   = EventoForm(request.POST or None, instance=evento)
    if form.is_valid():
        form.save()
        messages.success(request, 'Evento actualizado.')
        return redirect('agenda')
    return render(request, 'control_legal/registrar_evento.html', {
        'form': form, 'editando': True, 'evento': evento,
    })


# ── Agenda: eliminar evento ───────────────────────────────────
@login_required
def eliminar_evento(request, pk):
    evento = get_object_or_404(Evento, pk=pk)
    evento.delete()
    messages.success(request, 'Evento eliminado.')
    return redirect('agenda')

# ── Helper notificaciones ─────────────────────────────────────
def _notificar(destinatario, remitente, mensaje, url=''):
    if destinatario and destinatario != remitente:
        NotificacionInterna.objects.create(
            destinatario=destinatario,
            remitente=remitente,
            mensaje=mensaje,
            url_destino=url,
        )


# ── Tareas: lista ─────────────────────────────────────────────
@login_required
def lista_tareas(request):
    user = request.user
    if hasattr(user, 'perfil') and user.perfil.es_pasante():
        tareas = Tarea.objects.filter(asignada_a=user)
    else:
        tareas = Tarea.objects.all()

    estado_f    = request.GET.get('estado', '')
    prioridad_f = request.GET.get('prioridad', '')
    if estado_f:
        tareas = tareas.filter(estado=estado_f)
    if prioridad_f:
        tareas = tareas.filter(prioridad=prioridad_f)

    tareas = tareas.select_related('asignada_a', 'expediente__cliente').order_by('fecha_limite')
    return render(request, 'control_legal/lista_tareas.html', {
        'tareas':      tareas,
        'hoy':         timezone.now().date(),
        'estado_f':    estado_f,
        'prioridad_f': prioridad_f,
        'kanban_cols': [
            ('pendiente',  'Pendiente',  '#6c757d'),
            ('en_proceso', 'En proceso', '#0d6efd'),
            ('completada', 'Completada', '#198754'),
        ],
    })


# ── Tareas: crear ─────────────────────────────────────────────
@login_required
def crear_tarea(request):
    if hasattr(request.user, 'perfil') and request.user.perfil.es_pasante():
        messages.error(request, 'No tienes permiso para crear tareas.')
        return redirect('lista_tareas')
    form = TareaForm(request.POST or None)
    if form.is_valid():
        tarea            = form.save(commit=False)
        tarea.creada_por = request.user
        tarea.save()
        _notificar(
            tarea.asignada_a, request.user,
            f'Se te asignó la tarea "{tarea.titulo}" — vence {tarea.fecha_limite.strftime("%d/%m/%Y")}',
            f'/tareas/{tarea.pk}/',
        )
        messages.success(request, 'Tarea asignada correctamente.')
        return redirect('lista_tareas')
    return render(request, 'control_legal/registrar_tarea.html', {'form': form})


# ── Tareas: editar ────────────────────────────────────────────
@login_required
def editar_tarea(request, pk):
    tarea = get_object_or_404(Tarea, pk=pk)
    if hasattr(request.user, 'perfil') and request.user.perfil.es_pasante():
        messages.error(request, 'No tienes permiso para editar tareas.')
        return redirect('lista_tareas')
    form = TareaForm(request.POST or None, instance=tarea)
    if form.is_valid():
        form.save()
        messages.success(request, 'Tarea actualizada.')
        return redirect('lista_tareas')
    return render(request, 'control_legal/registrar_tarea.html', {
        'form': form, 'editando': True, 'tarea': tarea,
    })


# ── Tareas: detalle + evidencia + comentarios ─────────────────
@login_required
def detalle_tarea(request, pk):
    tarea    = get_object_or_404(Tarea, pk=pk)
    ev_form  = EvidenciaForm(None, None)
    com_form = ComentarioTareaForm()

    if request.method == 'POST':
        if 'subir_evidencia' in request.POST:
            ev_form = EvidenciaForm(request.POST, request.FILES)
            if ev_form.is_valid():
                ev            = ev_form.save(commit=False)
                ev.tarea      = tarea
                ev.subida_por = request.user
                ev.save()
                if tarea.estado == Tarea.ESTADO_PENDIENTE:
                    tarea.estado = Tarea.ESTADO_EN_PROCESO
                    tarea.save()
                if tarea.creada_por:
                    _notificar(
                        tarea.creada_por, request.user,
                        f'{request.user.get_full_name() or request.user.username} subió evidencia en "{tarea.titulo}"',
                        f'/tareas/{pk}/',
                    )
                messages.success(request, 'Evidencia subida correctamente.')
                return redirect('detalle_tarea', pk=pk)

        elif 'agregar_comentario' in request.POST:
            com_form = ComentarioTareaForm(request.POST)
            if com_form.is_valid():
                com       = com_form.save(commit=False)
                com.tarea = tarea
                com.autor = request.user
                com.save()
                receptor = tarea.creada_por if request.user == tarea.asignada_a else tarea.asignada_a
                _notificar(receptor, request.user, f'Nuevo comentario en "{tarea.titulo}"', f'/tareas/{pk}/')
                messages.success(request, 'Comentario agregado.')
                return redirect('detalle_tarea', pk=pk)

    puede_completar = not (hasattr(request.user, 'perfil') and request.user.perfil.es_pasante()) \
                      or request.user == tarea.asignada_a

    return render(request, 'control_legal/detalle_tarea.html', {
        'tarea':          tarea,
        'evidencias':     tarea.evidencias.all(),
        'comentarios':    tarea.comentarios.select_related('autor').all(),
        'ev_form':        ev_form,
        'com_form':       com_form,
        'puede_completar': puede_completar,
    })


# ── Tareas: cambiar estado ────────────────────────────────────
@login_required
def cambiar_estado_tarea(request, pk):
    tarea        = get_object_or_404(Tarea, pk=pk)
    nuevo_estado = request.POST.get('estado')
    es_pasante   = hasattr(request.user, 'perfil') and request.user.perfil.es_pasante()

    if es_pasante and nuevo_estado == Tarea.ESTADO_COMPLETADA:
        messages.error(request, 'Solo el abogado responsable puede marcar la tarea como completada.')
        return redirect('lista_tareas')

    if nuevo_estado in dict(Tarea.ESTADOS):
        tarea.estado = nuevo_estado
        if nuevo_estado == Tarea.ESTADO_COMPLETADA:
            tarea.fecha_completada = timezone.now()
            if tarea.es_recurrente:
                nueva = tarea.generar_siguiente_recurrente()
                if nueva:
                    messages.info(request, f'Tarea recurrente: nueva instancia creada con vencimiento {nueva.fecha_limite.strftime("%d/%m/%Y")}.')
            if tarea.creada_por:
                _notificar(
                    tarea.creada_por, request.user,
                    f'"{tarea.titulo}" fue completada por {request.user.get_full_name() or request.user.username}',
                    f'/tareas/{pk}/',
                )
        tarea.save()
        messages.success(request, f'Tarea marcada como {tarea.get_estado_display()}.')
    return redirect('lista_tareas')


# ── Notificaciones: marcar leídas ─────────────────────────────
@login_required
def marcar_notificaciones_leidas(request):
    NotificacionInterna.objects.filter(destinatario=request.user, leida=False).update(leida=True)
    return redirect(request.POST.get('next', 'dashboard'))


# ── Notificaciones: API JSON ──────────────────────────────────
@login_required
def api_notificaciones(request):
    qs = NotificacionInterna.objects.filter(
        destinatario=request.user, leida=False
    ).order_by('-creada_en')[:10]
    data = [
        {
            'id':        n.pk,
            'mensaje':   n.mensaje,
            'url':       n.url_destino,
            'creada_en': n.creada_en.strftime('%d/%m %H:%M'),
            'remitente': n.remitente.get_full_name() if n.remitente else '—',
        }
        for n in qs
    ]
    return JsonResponse({'count': len(data), 'notificaciones': data})

# ── Contratos: lista ──────────────────────────────────────────
@login_required
def lista_contratos(request):
    contratos = Contrato.objects.select_related(
        'cliente', 'generado_por').order_by('-fecha_generado')
    return render(request, 'control_legal/lista_contratos.html', {'contratos': contratos})


# ── Contratos: generar con plantilla o manual ─────────────────
@login_required
def generar_contrato(request):
    from django.utils import timezone
    form = GenerarContratoForm(request.POST or None)

    if form.is_valid():
        contrato             = form.save(commit=False)
        contrato.generado_por = request.user
        plantilla = form.cleaned_data.get('plantilla')
        if plantilla and not contrato.contenido_final.strip():
            cliente    = form.cleaned_data['cliente']
            expediente = form.cleaned_data.get('expediente')
            contenido  = plantilla.contenido
            contenido  = contenido.replace('{{cliente_nombre}}',    cliente.nombre_completo)
            contenido  = contenido.replace('{{cliente_ci}}',        cliente.ci)
            contenido  = contenido.replace('{{cliente_direccion}}', cliente.direccion or '')
            contenido  = contenido.replace('{{abogado_nombre}}',    request.user.get_full_name() or request.user.username)
            contenido  = contenido.replace('{{fecha_hoy}}',         timezone.now().strftime('%d de %B de %Y'))
            contenido  = contenido.replace('{{expediente_nurej}}',  expediente.nurej if expediente else '')
            contenido  = contenido.replace('{{monto}}',             str(contrato.monto or ''))
            contenido  = contenido.replace('{{ciudad}}',            contrato.ciudad)
            contrato.contenido_final = contenido
        contrato.save()
        messages.success(request, 'Contrato generado correctamente.')
        return redirect('ver_contrato', pk=contrato.pk)

    plantillas_json = json.dumps({
        str(p.pk): p.contenido
        for p in PlantillaContrato.objects.filter(activa=True)
    })
    return render(request, 'control_legal/generar_contrato.html', {
        'form': form,
        'plantillas_json': plantillas_json,
    })

# ── Contratos: ver ────────────────────────────────────────────
@login_required
def ver_contrato(request, pk):
    contrato = get_object_or_404(Contrato, pk=pk)
    return render(request, 'control_legal/ver_contrato.html', {'contrato': contrato})


# ── Contratos: exportar PDF ───────────────────────────────────
@login_required
def exportar_contrato_pdf(request, pk):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

    contrato     = get_object_or_404(Contrato, pk=pk)
    buf          = io.BytesIO()
    doc          = SimpleDocTemplate(buf, pagesize=letter,
                                     topMargin=3*cm, bottomMargin=2.5*cm,
                                     leftMargin=3*cm, rightMargin=3*cm)
    styles       = getSampleStyleSheet()
    titulo_style = ParagraphStyle('titulo', parent=styles['Normal'],
                                  fontSize=13, fontName='Helvetica-Bold',
                                  alignment=TA_CENTER, spaceAfter=20)
    cuerpo_style = ParagraphStyle('cuerpo', parent=styles['Normal'],
                                  fontSize=11, leading=18,
                                  alignment=TA_JUSTIFY, spaceAfter=12)
    firma_style  = ParagraphStyle('firma', parent=styles['Normal'],
                                  fontSize=10, spaceBefore=40)
    elements = []

    elements.append(Paragraph("LexNova Bolivia", ParagraphStyle(
        'header', parent=styles['Normal'], fontSize=10,
        fontName='Helvetica', alignment=TA_CENTER)))
    elements.append(Spacer(1, 0.3*cm))
    elements.append(Paragraph(contrato.titulo.upper(), titulo_style))
    elements.append(Spacer(1, 0.5*cm))

    for parrafo in contrato.contenido_final.split('\n'):
        if parrafo.strip():
            elements.append(Paragraph(parrafo.strip(), cuerpo_style))
        else:
            elements.append(Spacer(1, 0.3*cm))

    elements.append(Spacer(1, 2*cm))
    elements.append(Paragraph(
        f"________________________________&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
        f"________________________________", firma_style))
    elements.append(Paragraph(
        f"{contrato.cliente.nombre_completo}&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
        f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
        f"{contrato.generado_por.get_full_name() or 'Abogado'}",
        ParagraphStyle('firma_nombre', parent=styles['Normal'], fontSize=9)))

    doc.build(elements)
    buf.seek(0)
    nombre = f"Contrato_{contrato.cliente.nombre_completo.replace(' ', '_')}.pdf"
    return HttpResponse(buf, content_type='application/pdf',
                        headers={'Content-Disposition': f'inline; filename="{nombre}"'})


# ── Plantillas: lista ─────────────────────────────────────────
@login_required
def lista_plantillas(request):
    plantillas = PlantillaContrato.objects.filter(activa=True).order_by('nombre')
    return render(request, 'control_legal/lista_plantillas.html', {'plantillas': plantillas})


# ── Plantillas: crear ─────────────────────────────────────────
@login_required
def crear_plantilla(request):
    form = PlantillaContratoForm(request.POST or None)
    if form.is_valid():
        p            = form.save(commit=False)
        p.creada_por = request.user
        p.save()
        messages.success(request, 'Plantilla creada correctamente.')
        return redirect('lista_plantillas')
    return render(request, 'control_legal/crear_plantilla.html', {'form': form})


# ── Plantillas: editar ────────────────────────────────────────
@login_required
def editar_plantilla(request, pk):
    plantilla = get_object_or_404(PlantillaContrato, pk=pk)
    form      = PlantillaContratoForm(request.POST or None, instance=plantilla)
    if form.is_valid():
        form.save()
        messages.success(request, 'Plantilla actualizada.')
        return redirect('lista_plantillas')
    return render(request, 'control_legal/crear_plantilla.html', {
        'form': form, 'editando': True, 'plantilla': plantilla,
    })

import google.generativeai as genai
from django.conf import settings

SYSTEM_PROMPT_CONTRATOS = """
Actúa como un Abogado Senior y Asesor Legal Corporativo en Bolivia, con más de 20 años de experiencia en Derecho Civil, Comercial y Tecnológico. Tu objetivo es redactar contratos jurídicamente inexpugnables, precisos, con visión preventiva de litigios y estrictamente adaptados a la normativa del Estado Plurinacional de Bolivia.

SISTEMA HÍBRIDO (AUTOMATIZACIÓN + EDICIÓN MANUAL)
Tu redacción debe estar diseñada para ser un documento "llave en mano", pero con espacios de edición manual a prueba de errores. Para todo dato faltante, variable, fecha o monto, utiliza CORCHETES Y MAYÚSCULAS (ej. [NOMBRE COMPLETO DEL COMPRADOR], [NÚMERO DE C.I. Y EXPEDICIÓN], [MONTO EN NÚMEROS] ([MONTO EN LETRAS] 00/100 BOLIVIANOS)). Nunca asumas ni inventes datos de las partes.

BASE DE CONOCIMIENTO Y MARCO NORMATIVO
- Derecho Civil (D.L. No 12760): contratos de transferencia, arrendamiento, anticresis, mutuo y garantías. Aplicación estricta de los 4 requisitos de validez del Art. 452.
- Derecho Comercial (Código de Comercio) y Corporativo.
- Contratos Tecnológicos: SaaS, IaaS, PaaS, desarrollo de software y licenciamiento con cláusulas de SLA, propiedad intelectual (SENAPI), confidencialidad (NDA), privacidad de datos y limitación de responsabilidad técnica.

INSTRUCCIONES DE ESTILO BOLIVIANO
- Usa terminología notarial boliviana: "mayores de edad y hábiles por derecho", "sin que medie dolo, presión, fraude o vicio alguno en el consentimiento", "a su entera satisfacción".
- En transferencias incluye: "alodial y libre de todo gravamen" y cláusula de "evicción y saneamiento conforme a ley".
- Enumera cláusulas con ordinales en mayúsculas: PRIMERA.- (DE LAS PARTES Y SU DERECHO PROPIETARIO).
- Incluye siempre: Resolución de Pleno Derecho (Art. 569 C.C.), cláusulas penales por mora y jurisdicción competente.
- Contratos civiles de baja cuantía: vía ordinaria. Contratos comerciales/tecnológicos: arbitraje institucional (Ley 708), CNC o CAINCO.
- Si requiere Derechos Reales: formato MINUTA. Si es privado: cierra con "al solo reconocimiento de firmas y rúbricas ante autoridad competente, surtirá los efectos de instrumento público".

ESTRUCTURA ESTÁNDAR
1. Encabezado con identificación plena de partes.
2. PRIMERA.- (ANTECEDENTES / NATURALEZA)
3. SEGUNDA.- (OBJETO DEL CONTRATO)
4. Cláusulas operativas: Precio, Pago, Plazos, Obligaciones.
5. Cláusulas de salvaguarda: Prohibiciones, Confidencialidad.
6. Cláusula de Incumplimiento, Penalidades y Resolución.
7. Cláusula de Jurisdicción y Competencia.
8. Cláusula de Aceptación y Conformidad con espacio para firmas.

Si el usuario omite detalles vitales, usa estándares lógicos entre corchetes para revisión.
Genera el contrato COMPLETO, listo para copiar en Word, rellenar y firmar.
Responde SOLO con el texto del contrato, sin explicaciones adicionales ni markdown.
"""


@login_required
def generar_contrato_ia(request):
    contrato_generado = None
    error = None

    if request.method == 'POST':
        tipo        = request.POST.get('tipo_contrato', '')
        partes      = request.POST.get('partes', '')
        objeto      = request.POST.get('objeto', '')
        precio      = request.POST.get('precio', '')
        plazo       = request.POST.get('plazo', '')
        condiciones = request.POST.get('condiciones', '')

        prompt_completo = f"""{SYSTEM_PROMPT_CONTRATOS}

[DATOS DEL CONTRATO A GENERAR]
Tipo de Contrato: {tipo}
Identificación de las Partes: {partes}
Objeto y Detalles Específicos: {objeto}
Precio / Contraprestación: {precio}
Plazo de vigencia: {plazo}
Condiciones especiales adicionales: {condiciones}

Redacta el contrato completo según las instrucciones anteriores.
"""

        try:
            genai.configure(api_key=settings.GEMINI_API_KEY)
            modelo   = genai.GenerativeModel('gemini-2.5-flash')
            response = modelo.generate_content(prompt_completo)
            contrato_generado = response.text

            if request.POST.get('guardar') and contrato_generado:
                cliente_id = request.POST.get('cliente_id')
                cliente    = Cliente.objects.filter(pk=cliente_id).first()
                if cliente:
                    Contrato.objects.create(
                        titulo          = f"{tipo} — {cliente.nombre_completo}",
                        cliente         = cliente,
                        contenido_final = contrato_generado,
                        generado_por    = request.user,
                        ciudad          = "La Paz",
                    )
                    messages.success(request, 'Contrato guardado correctamente.')
                    return redirect('lista_contratos')

        except Exception as e:
            error = f"Error al conectar con Gemini: {str(e)}"

    clientes = Cliente.objects.filter(activo=True).order_by('nombre_completo')
    return render(request, 'control_legal/generar_contrato_ia.html', {
        'contrato_generado': contrato_generado,
        'clientes':          clientes,
        'error':             error,
        'post':              request.POST,
    })


# ── Reportes con gráficos ─────────────────────────────────────
@login_required
def reportes(request):
    from django.db.models import Count
    from django.db.models.functions import TruncMonth
    from django.utils import timezone
    from datetime import timedelta
    import json

    hoy  = timezone.now().date()
    anio = hoy.year

    # ── Expedientes por mes (últimos 12 meses) ────────────────
    # CORRECCIÓN: fecha_inicio es DateField, TruncMonth funciona
    # pero con USE_TZ=True puede dar None en el strftime si el queryset está vacío.
    # Se protege con un condicional en la list comprehension.
    hace_12 = hoy - timedelta(days=365)
    exp_por_mes = (
        Expediente.objects
        .filter(fecha_inicio__gte=hace_12)
        .annotate(mes=TruncMonth('fecha_inicio'))
        .values('mes')
        .annotate(total=Count('id'))
        .order_by('mes')
    )
    # CORRECCIÓN BUG: se verifica que e['mes'] no sea None antes de formatear
    labels_meses = [e['mes'].strftime('%b %Y') for e in exp_por_mes if e['mes'] is not None]
    data_meses   = [e['total'] for e in exp_por_mes if e['mes'] is not None]

    # ── Expedientes por materia ───────────────────────────────
    exp_por_materia = (
        Expediente.objects
        .filter(activo=True)
        .values('materia')
        .annotate(total=Count('id'))
        .order_by('-total')
    )
    labels_materias = [e['materia'] for e in exp_por_materia]
    data_materias   = [e['total'] for e in exp_por_materia]

    # ── Expedientes por estado ────────────────────────────────
    exp_por_estado = (
        Expediente.objects
        .values('estado')
        .annotate(total=Count('id'))
    )
    estados_map = {'activo': 'Activo', 'concluido': 'Concluido', 'archivado': 'Archivado'}
    labels_estados = [estados_map.get(e['estado'], e['estado']) for e in exp_por_estado]
    data_estados   = [e['total'] for e in exp_por_estado]

    # ── Expedientes por abogado ───────────────────────────────
    exp_por_abogado = (
        Expediente.objects
        .filter(activo=True, abogado_asignado__isnull=False)
        .values('abogado_asignado__first_name', 'abogado_asignado__last_name')
        .annotate(total=Count('id'))
        .order_by('-total')[:8]
    )
    labels_abogados = [
        f"{e['abogado_asignado__first_name']} {e['abogado_asignado__last_name']}".strip() or 'Sin nombre'
        for e in exp_por_abogado
    ]
    data_abogados = [e['total'] for e in exp_por_abogado]

    # ── Tareas por estado ─────────────────────────────────────
    tareas_estado = (
        Tarea.objects
        .values('estado')
        .annotate(total=Count('id'))
    )
    t_map = {'pendiente': 'Pendiente', 'en_proceso': 'En proceso', 'completada': 'Completada'}
    labels_tareas = [t_map.get(t['estado'], t['estado']) for t in tareas_estado]
    data_tareas   = [t['total'] for t in tareas_estado]

    # ── Resumen general ───────────────────────────────────────
    resumen = {
        'total_clientes':    Cliente.objects.filter(activo=True).count(),
        'total_expedientes': Expediente.objects.count(),
        'activos':           Expediente.objects.filter(estado='activo').count(),
        'concluidos':        Expediente.objects.filter(estado='concluido').count(),
        'total_contratos':   Contrato.objects.count(),
        'tareas_pendientes': Tarea.objects.filter(estado='pendiente').count(),
        'tareas_vencidas':   Tarea.objects.filter(
                                estado__in=['pendiente', 'en_proceso'],
                                fecha_limite__lt=hoy).count(),
        'eventos_mes':       Evento.objects.filter(
                                fecha_hora__year=anio,
                                fecha_hora__month=hoy.month).count(),
    }

    return render(request, 'control_legal/reportes.html', {
        'resumen':          resumen,
        'labels_meses':     json.dumps(labels_meses),
        'data_meses':       json.dumps(data_meses),
        'labels_materias':  json.dumps(labels_materias),
        'data_materias':    json.dumps(data_materias),
        'labels_estados':   json.dumps(labels_estados),
        'data_estados':     json.dumps(data_estados),
        'labels_abogados':  json.dumps(labels_abogados),
        'data_abogados':    json.dumps(data_abogados),
        'labels_tareas':    json.dumps(labels_tareas),
        'data_tareas':      json.dumps(data_tareas),
    })


# ── Portal del cliente ────────────────────────────────────────
# Estas vistas son PÚBLICAS — no requieren login del bufete

def portal_login(request):
    """Página donde el cliente ingresa su código de acceso."""
    error = None
    if request.method == 'POST':
        codigo = request.POST.get('codigo', '').strip().upper()
        try:
            cliente = Cliente.objects.get(codigo_acceso=codigo, activo=True)
            request.session['portal_cliente_id'] = cliente.pk
            return redirect('portal_cliente')
        except Cliente.DoesNotExist:
            error = 'Código incorrecto. Verifica el código que te proporcionó tu abogado.'
    return render(request, 'control_legal/portal_login.html', {'error': error})


def portal_cliente(request):
    """Panel del cliente — muestra sus expedientes, documentos y movimientos."""
    cliente_id = request.session.get('portal_cliente_id')
    if not cliente_id:
        return redirect('portal_login')

    try:
        cliente = Cliente.objects.get(pk=cliente_id, activo=True)
    except Cliente.DoesNotExist:
        del request.session['portal_cliente_id']
        return redirect('portal_login')

    expedientes = cliente.expedientes.filter(activo=True).prefetch_related(
        'documentos', 'movimientos'
    ).order_by('-fecha_inicio')

    return render(request, 'control_legal/portal_cliente.html', {
        'cliente':     cliente,
        'expedientes': expedientes,
    })


def portal_logout(request):
    """Cierra la sesión del portal del cliente."""
    request.session.flush()
    return redirect('portal_login')


SYSTEM_PROMPT_MEMORIALES = """
Actúa como un Abogado Litigante Senior en Bolivia, especializado en redacción de escritos judiciales ante el Órgano Judicial del Estado Plurinacional de Bolivia.

Tu tarea es redactar memoriales judiciales completos, formalmente correctos y adaptados al sistema procesal boliviano.

REGLAS DE REDACCIÓN
- Usa el encabezado formal: "SEÑOR JUEZ [materia] DE PARTIDO / SEÑOR JUEZ DE INSTRUCCIÓN EN LO [materia]" según corresponda.
- Identifica al abogado con su nombre y matrícula: "bajo patrocinio del Abogado [nombre], con Matrícula del Colegio de Abogados N° [MATRÍCULA]".
- Usa fórmulas procesales bolivianas: "A Ud. con el debido respeto me dirijo exponiendo", "POR TANTO", "A USTED SEÑOR JUEZ pido se sirva...".
- Fundamenta cada petición con el artículo legal correspondiente.
- Estructura con secciones en mayúsculas: ANTECEDENTES, FUNDAMENTOS DE DERECHO, PETITORIO.
- El PETITORIO debe ser claro, numerado y específico.
- Para datos faltantes usa CORCHETES: [MATRÍCULA DEL ABOGADO], [NÚMERO DE EXPEDIENTE], [FECHA DE AUDIENCIA].
- Cierra con: "Es justicia que espero merecer. [Ciudad], [fecha]."

ESTILO
- Formal, preciso, sin ambigüedades.
- Párrafos cortos y numerados donde corresponda.
- Responde SOLO con el texto del memorial, sin explicaciones ni markdown.
"""


@login_required
def lista_memoriales(request):
    memoriales = Memorial.objects.select_related(
        'cliente', 'plantilla', 'generado_por'
    ).order_by('-fecha_generado')
    return render(request, 'control_legal/lista_memoriales.html', {'memoriales': memoriales})


@login_required
def generar_memorial(request):
    from django.utils import timezone

    plantillas  = PlantillaMemorial.objects.filter(activa=True).order_by('categoria', 'nombre')
    clientes    = Cliente.objects.filter(activo=True).order_by('nombre_completo')
    expedientes = Expediente.objects.filter(activo=True).select_related('cliente').order_by('-id')

    memorial_generado = None
    error             = None
    plantilla_sel     = None

    if request.method == 'POST':
        plantilla_id  = request.POST.get('plantilla_id')
        cliente_id    = request.POST.get('cliente_id')
        expediente_id = request.POST.get('expediente_id')
        hechos        = request.POST.get('hechos', '')
        peticion      = request.POST.get('peticion', '')
        datos_extra   = request.POST.get('datos_extra', '')

        plantilla_sel = PlantillaMemorial.objects.filter(pk=plantilla_id).first()
        cliente       = Cliente.objects.filter(pk=cliente_id).first()
        expediente    = Expediente.objects.filter(pk=expediente_id).first() if expediente_id else None

        if not plantilla_sel or not cliente:
            error = 'Debes seleccionar una plantilla y un cliente.'
        else:
            contexto_variables = f"""
Cliente: {cliente.nombre_completo}
CI del cliente: {cliente.ci}
Dirección: {cliente.direccion or '[DIRECCIÓN]'}
Abogado: {request.user.get_full_name() or '[NOMBRE DEL ABOGADO]'}
Ciudad: La Paz
Fecha: {timezone.now().strftime('%d de %B de %Y')}
NUREJ/Expediente: {expediente.nurej if expediente else '[NUREJ]'}
Juzgado: {expediente.juzgado if expediente else '[JUZGADO]'}
Materia: {expediente.materia if expediente else '[MATERIA]'}
"""
            prompt = f"""{SYSTEM_PROMPT_MEMORIALES}

PLANTILLA A USAR — {plantilla_sel.nombre}:
{plantilla_sel.estructura}

NORMAS APLICABLES:
{plantilla_sel.normas_aplicables or 'Aplica las normas que correspondan según la materia.'}

DATOS DEL CASO:
{contexto_variables}

HECHOS DEL CASO (proporcionados por el abogado):
{hechos}

PETICIÓN ESPECÍFICA:
{peticion}

DATOS ADICIONALES:
{datos_extra or 'Ninguno.'}

Redacta el memorial completo siguiendo la plantilla y adaptándola al caso concreto.
"""
            try:
                genai.configure(api_key=settings.GEMINI_API_KEY)
                modelo   = genai.GenerativeModel('gemini-2.5-flash')
                response = modelo.generate_content(prompt)
                memorial_generado = response.text

                if request.POST.get('guardar') and memorial_generado and cliente:
                    Memorial.objects.create(
                        plantilla       = plantilla_sel,
                        expediente      = expediente,
                        cliente         = cliente,
                        generado_por    = request.user,
                        titulo          = f"{plantilla_sel.nombre} — {cliente.nombre_completo}",
                        contenido_final = memorial_generado,
                    )
                    messages.success(request, 'Memorial guardado correctamente.')
                    return redirect('lista_memoriales')

            except Exception as e:
                error = f"Error al conectar con Gemini: {str(e)}"

    return render(request, 'control_legal/generar_memorial.html', {
        'plantillas':        plantillas,
        'clientes':          clientes,
        'expedientes':       expedientes,
        'memorial_generado': memorial_generado,
        'plantilla_sel':     plantilla_sel,
        'error':             error,
        'post':              request.POST,
    })


@login_required
def ver_memorial(request, pk):
    memorial = get_object_or_404(Memorial, pk=pk)
    return render(request, 'control_legal/ver_memorial.html', {'memorial': memorial})


@login_required
def exportar_memorial_pdf(request, pk):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

    memorial     = get_object_or_404(Memorial, pk=pk)
    buf          = io.BytesIO()
    doc          = SimpleDocTemplate(buf, pagesize=letter,
                                     topMargin=3*cm, bottomMargin=2.5*cm,
                                     leftMargin=3*cm, rightMargin=3*cm)
    styles       = getSampleStyleSheet()
    titulo_style = ParagraphStyle('titulo', parent=styles['Normal'],
                                  fontSize=12, fontName='Helvetica-Bold',
                                  alignment=TA_CENTER, spaceAfter=16)
    cuerpo_style = ParagraphStyle('cuerpo', parent=styles['Normal'],
                                  fontSize=11, leading=18,
                                  alignment=TA_JUSTIFY, spaceAfter=10)
    elements     = []

    elements.append(Paragraph(memorial.titulo.upper(), titulo_style))
    elements.append(Spacer(1, 0.5*cm))

    for parrafo in memorial.contenido_final.split('\n'):
        if parrafo.strip():
            elements.append(Paragraph(parrafo.strip(), cuerpo_style))
        else:
            elements.append(Spacer(1, 0.25*cm))

    doc.build(elements)
    buf.seek(0)
    nombre = f"Memorial_{memorial.cliente.nombre_completo.replace(' ','_')}.pdf"
    return HttpResponse(buf, content_type='application/pdf',
                        headers={'Content-Disposition': f'inline; filename="{nombre}"'})


# =====================================================================
# CALCULADORA DE ASISTENCIA FAMILIAR — Ley 603
# Correcciones aplicadas:
#   1. context siempre recibe 'hoy' y los campos del formulario
#      para que NO se borren al calcular (Bug repoblación)
#   2. distrito_actual / juzgado_actual / etc. se setean SIEMPRE
#      desde POST o desde valores_iniciales (Bug campos vacíos)
#   3. session.modified = True fuerza el guardado antes del redirect
#      al PDF (Bug PDF en blanco)
# =====================================================================
@login_required
def calculadora_asistencia(request, caso_id=None):
    expediente = None
    if caso_id:
        expediente = get_object_or_404(Expediente, pk=caso_id)

    hoy_date = timezone.now().date()

    valores_iniciales = {
        'distrito_judicial': 'La Paz',
        'juzgado':           expediente.juzgado if expediente else '',
        'nurej_ianus':       expediente.nurej    if expediente else '',
        'partes':            str(expediente.cliente) if expediente else '',
        'beneficiarios':     '',
        'nro_resolucion':    '',
        'monto_mensual':     '',
        'fecha_inicio':      '',
    }

    context = {
        'expediente':            expediente,
        'valores_iniciales':     valores_iniciales,
        'distritos':             asistencia_services.DISTRITOS_JUDICIALES,
        'resultado':             None,
        'error':                 None,
        'hoy':                   hoy_date.strftime('%Y-%m-%d'),
        'distrito_actual':       valores_iniciales['distrito_judicial'],
        'juzgado_actual':        valores_iniciales['juzgado'],
        'nurej_actual':          valores_iniciales['nurej_ianus'],
        'partes_actual':         valores_iniciales['partes'],
        'beneficiarios_actual':  '',
        'nro_resolucion_actual': '',
        'aplicar_duodecimas':    False,
    }

    if request.method == 'POST':
        # Repoblar campos SIEMPRE para que no se borren
        context['distrito_actual']       = request.POST.get('distrito_judicial', valores_iniciales['distrito_judicial'])
        context['juzgado_actual']        = request.POST.get('juzgado', '')
        context['nurej_actual']          = request.POST.get('nurej_ianus', '')
        context['partes_actual']         = request.POST.get('partes', '')
        context['beneficiarios_actual']  = request.POST.get('beneficiarios', '')
        context['nro_resolucion_actual'] = request.POST.get('nro_resolucion', '')
        context['post']                  = request.POST
        # Repoblar el tramo 0 para que no se borre al calcular
        context['tramo_inicio_0'] = request.POST.get('tramo_inicio_0', '')
        context['tramo_fin_0']    = request.POST.get('tramo_fin_0', hoy_date.strftime('%Y-%m-%d'))
        context['tramo_monto_0']  = request.POST.get('tramo_monto_0', '')

        accion             = request.POST.get('accion', 'calcular')
        aplicar_duodecimas = request.POST.get('aplicar_duodecimas') == 'on'
        context['aplicar_duodecimas'] = aplicar_duodecimas

        tramos = asistencia_services.parsear_tramos_desde_post(request.POST)
        pagos  = asistencia_services.parsear_pagos_desde_post(request.POST)

        if not tramos:
            context['error'] = (
                'No se recibieron tramos válidos. '
                'Asegúrate de ingresar al menos un tramo con fechas y monto correctos.'
            )
            return render(request, 'control_legal/calculadora_asistencia.html', context)

        try:
            resultado_calculo = asistencia_services.computar_liquidacion_completa(
                tramos,
                aplicar_duodecimas=aplicar_duodecimas,
                pagos=pagos,
            )

            # ── REPORTE PDF: generar y devolver directamente ──────
            # NO usamos redirect ni sesión — generamos el PDF aquí
            # mismo con los datos que ya tenemos en memoria.
            if accion == 'reporte_pdf':
                datos_planilla = {
                    'distrito_judicial':  request.POST.get('distrito_judicial', ''),
                    'juzgado':            request.POST.get('juzgado', ''),
                    'nurej_ianus':        request.POST.get('nurej_ianus', ''),
                    'partes':             request.POST.get('partes', ''),
                    'beneficiarios':      request.POST.get('beneficiarios', ''),
                    'nro_resolucion':     request.POST.get('nro_resolucion', ''),
                    'aplicar_duodecimas': aplicar_duodecimas,
                    'total_liquidacion':  resultado_calculo['total_liquidacion'],
                    'total_pagado':       resultado_calculo.get('total_pagado', Decimal('0')),
                    'saldo_adeudado':     resultado_calculo.get('saldo_adeudado', resultado_calculo['total_liquidacion']),
                    'detalle_pagos':      resultado_calculo.get('detalle_pagos', []),
                    'filas':              resultado_calculo['filas'],
                }
                buffer_pdf = asistencia_services.generar_pdf_liquidacion_asistencia(datos_planilla)
                response = HttpResponse(buffer_pdf.read(), content_type='application/pdf')
                response['Content-Disposition'] = 'inline; filename="Planilla_Liquidacion_Ley603.pdf"'
                return response

            # ── CALCULAR: mostrar resultado en pantalla ───────────
            context['resultado'] = resultado_calculo

        except Exception as e:
            context['error'] = f'Error al calcular la liquidación: {str(e)}'

    return render(request, 'control_legal/calculadora_asistencia.html', context)

# ── Calculadora integrada al expediente ───────────────────────
@login_required
def calculadora_asistencia_expediente(request, pk):
    return calculadora_asistencia(request, caso_id=pk)

# ── Endpoint IA para extracción de datos del expediente ───────
@login_required
def api_extraer_ia_liquidacion(request, pk):
    """
    Extrae datos de los documentos del expediente usando Gemini.
    CORRECCIÓN: Ahora responde con 'ok' (bool) en lugar de 'status'
    para que el JS del template pueda leer correctamente el resultado.
    """
    expediente = get_object_or_404(Expediente, pk=pk)
    datos, error = asistencia_services.extraer_datos_liquidacion_ia(expediente)
    if error:
        return JsonResponse({'ok': False, 'error': error})
    return JsonResponse({'ok': True, 'datos': datos})


# ── Generación y descarga del PDF de liquidación ──────────────
@login_required
def exportar_liquidacion_pdf(request):
    """
    Recupera los cálculos de la sesión y genera el PDF judicial.
    """
    pdf_raw = request.session.get('pdf_data')
    if not pdf_raw:
        return HttpResponse(
            "No se han encontrado cálculos activos. Realice el cálculo primero.",
            status=400,
        )

    datos_planilla = {
        'distrito_judicial':  pdf_raw['distrito_judicial'],
        'juzgado':            pdf_raw['juzgado'],
        'nurej_ianus':        pdf_raw['nurej_ianus'],
        'partes':             pdf_raw['partes'],
        'aplicar_duodecimas': pdf_raw['aplicar_duodecimas'],
        'total_liquidacion':  Decimal(str(pdf_raw['total_liquidacion'])),
        'filas': [
            {
                'monto_acordado': Decimal(str(f['monto_acordado'])),
                'fecha_del':      datetime.strptime(f['fecha_del'], '%Y-%m-%d').date(),
                'fecha_al':       datetime.strptime(f['fecha_al'],  '%Y-%m-%d').date(),
                'meses_label':    f['meses_label'],
                'monto_pagar':    Decimal(str(f['monto_pagar'])),
            } for f in pdf_raw['filas']
        ],
    }

    buffer_pdf = asistencia_services.generar_pdf_liquidacion_asistencia(datos_planilla)

    response = HttpResponse(buffer_pdf.read(), content_type='application/pdf')
    response['Content-Disposition'] = 'inline; filename="Planilla_Liquidacion_Ley603.pdf"'
    return response


"""
FRAGMENTO PARA views.py — línea 1155
=====================================
Reemplaza SOLO la función jurisprudencia_asistente (las 4 líneas vacías).
USA el services.py que ya existe en tu proyecto.

Añade este import al inicio de views.py si no está:
    from . import services as srv
"""

# ── Import a añadir al inicio de views.py (junto a los demás imports) ──
# from . import services as srv


@login_required
def jurisprudencia_asistente(request):
    """
    Centro de Jurisprudencia — conectado al services.py existente.
      A. Estrategia legal con IA (usa sugerir_estrategia_jurisprudencial)
      B. Banco de precedentes local (usa buscar_precedentes_locales)
      C. Bot buscador multi-fuente (usa extraer_sentencia_tcp con respaldo web)
      D. Analizar texto pegado de sentencia (usa Gemini directo)
    """
    from .models import PrecedenteJurisprudencial
    from . import services as srv

    materias = PrecedenteJurisprudencial.MATERIAS
    opcion   = request.POST.get('opcion') if request.method == 'POST' else None

    estrategia          = None
    precedente_extraido = None
    extraido_fuente_alt = False
    analisis_texto      = None
    resultados_busqueda = []
    error               = None

    # ── OPCIÓN A: Estrategia legal con IA ───────────────────
    if opcion == 'estrategia':
        materia  = request.POST.get('materia', '')
        hechos   = request.POST.get('hechos_caso', '')
        peticion = request.POST.get('peticion', '')

        if not hechos.strip():
            error = "Debes describir los hechos del caso."
        else:
            # Buscar precedentes locales para darle contexto a la IA
            precs_locales = srv.buscar_precedentes_locales(
                termino=request.POST.get('termino_precedentes', ''),
                materia=materia,
                limite=5,
            )
            contexto = srv.precedentes_a_contexto(precs_locales)

            estrategia, error = srv.sugerir_estrategia_jurisprudencial(
                hechos_caso=hechos,
                materia=materia,
                peticion=peticion,
                precedentes_contexto=contexto,
            )

    # ── OPCIÓN C: Bot multi-fuente (TCP + respaldo web) ─────
    elif opcion == 'extraccion_tcp':
        entrada = request.POST.get('url_tcp', '').strip()
        if not entrada:
            error = "Ingresa un número de sentencia o URL."
        else:
            datos, error = srv.extraer_sentencia_tcp(entrada)
            if datos and not error:
                extraido_fuente_alt = datos.get('fuente_alternativa', False)
                precedente_extraido, creado, err_guardar = srv.guardar_precedente(
                    datos, usuario=request.user
                )
                if err_guardar:
                    error = err_guardar
                elif creado:
                    messages.success(
                        request,
                        f'Sentencia {precedente_extraido.numero_sentencia} guardada en el banco.'
                    )
                else:
                    messages.info(
                        request,
                        f'La sentencia {precedente_extraido.numero_sentencia} ya existía en el banco.'
                    )

    # ── OPCIÓN D: Analizar texto pegado ────────────────────
    elif opcion == 'analizar_texto':
        texto_sent = request.POST.get('texto_sentencia', '').strip()
        if not texto_sent:
            error = "Pega el texto de la sentencia."
        else:
            from . import jurisprudencia_services as juri_srv
            analisis_texto, error = juri_srv.analizar_sentencia_con_ia(texto_sent)

    # ── OPCIÓN B-BUSCAR: Búsqueda libre con IA ─────────────
    elif opcion == 'busqueda_libre':
        termino_libre = request.POST.get('termino_libre', '').strip()
        materia_lib   = request.POST.get('materia_libre', '')
        if not termino_libre:
            error = "Ingresa un término de búsqueda."
        else:
            from . import jurisprudencia_services as juri_srv
            resultados_busqueda, error = juri_srv.buscar_precedentes_libres(
                termino_libre, materia_lib
            )

    # ── Banco de precedentes guardados (siempre visible) ───
    buscar         = request.GET.get('buscar', '')
    materia_filtro = request.GET.get('materia', '')

    banco = srv.buscar_precedentes_locales(
        termino=buscar,
        materia=materia_filtro,
        limite=50,
    )

    return render(request, 'control_legal/jurisprudencia_asistente.html', {
        'materias':                    materias,
        'estrategia':                  estrategia,
        'precedente_extraido':         precedente_extraido,
        'extraido_fuente_alternativa': extraido_fuente_alt,
        'analisis_texto':              analisis_texto,
        'resultados_busqueda':         resultados_busqueda,
        'precedentes_banco':           banco,
        'buscar':                      buscar,
        'materia_filtro':              materia_filtro,
        'error':                       error,
        'post':                        request.POST if request.method == 'POST' else {},
    })


@login_required
def guardar_precedente_desde_busqueda(request):
    """Guarda en BD un precedente encontrado por búsqueda libre."""
    from .models import PrecedenteJurisprudencial
    from . import services as srv
    from datetime import datetime

    if request.method != 'POST':
        return redirect('jurisprudencia_asistente')

    numero = request.POST.get('numero_sentencia', '').strip()
    if not numero:
        messages.error(request, 'Número de sentencia requerido.')
        return redirect('jurisprudencia_asistente')

    fecha_str = request.POST.get('fecha_resolucion', '')
    fecha_obj = None
    if fecha_str:
        try:
            fecha_obj = datetime.strptime(fecha_str, '%Y-%m-%d').date()
        except ValueError:
            pass

    datos = {
        'numero_sentencia':   numero,
        'tipo_resolucion':    request.POST.get('tipo_resolucion', 'otro'),
        'tribunal':           request.POST.get('tribunal', 'tcp'),
        'magistrado_relator': request.POST.get('magistrado_relator', ''),
        'fecha_resolucion':   fecha_obj,
        'materia':            request.POST.get('materia', 'constitucional'),
        'ratio_decidendi':    request.POST.get('ratio_decidendi', ''),
        'resumen_ia':         request.POST.get('resumen', ''),
        'palabras_clave':     request.POST.get('palabras_clave', ''),
        'url_fuente':         request.POST.get('url_fuente', ''),
        'origen':             PrecedenteJurisprudencial.ORIGEN_IA,
        'fuente_alternativa': True,
    }

    _, creado, err = srv.guardar_precedente(datos, usuario=request.user)
    if err:
        messages.error(request, err)
    elif creado:
        messages.success(request, f'Sentencia {numero} guardada en el banco.')
    else:
        messages.info(request, f'La sentencia {numero} ya existía en el banco.')

    return redirect('jurisprudencia_asistente')


@login_required
def eliminar_precedente(request, pk):
    """Elimina un precedente del banco."""
    from .models import PrecedenteJurisprudencial
    p = get_object_or_404(PrecedenteJurisprudencial, pk=pk)
    numero = p.numero_sentencia
    p.delete()
    messages.success(request, f'Precedente {numero} eliminado del banco.')
    return redirect('jurisprudencia_asistente')

# ══════════════════════════════════════════════════════════════
# CHATBOT JURÍDICO — Gemini 2.5 Flash
# Endpoint AJAX que recibe el historial completo de la
# conversación y devuelve la respuesta en JSON.
# ══════════════════════════════════════════════════════════════

SYSTEM_PROMPT_CHATBOT = """
Eres LexBot, el asistente jurídico inteligente del despacho LexNova, especializado \
en Derecho boliviano. Tienes conocimiento profundo y actualizado de:

LEGISLACIÓN PRINCIPAL
- Constitución Política del Estado Plurinacional de Bolivia (CPE)
- Código Civil (D.L. 12760) y Código de Comercio
- Código Procesal Civil (Ley 439) y Código Procesal Penal (Ley 1970)
- Código Penal (D.L. 10426) Eevado a rango de ley por la Ley 1768
- Ley N° 603 — Código de las Familias y del Proceso Familiar (CFPF)
- Ley N° 348 — Ley Integral para Garantizar a las Mujeres una Vida Libre de Violencia
- Ley N° 045 — Ley contra el Racismo y toda forma de Discriminación
- Ley N° 2028 — Ley de Municipalidades
- Ley N° 1178 — SAFCO
- Ley N° 004 — Ley Marcelo Quiroga Santa Cruz (anticorrupción)
- Ley N° 2341 — Ley de Procedimiento Administrativo
- Código Tributario (Ley 2492) y normativa SIN/ADUANA
- Ley General del Trabajo y DS 224 (reglamento)
- Ley N° 260 — Ley Orgánica del Ministerio Público
- Ley N° 025 — Ley del Órgano Judicial

ESPECIALIDADES
- Derecho de Familia: divorcio, asistencia familiar (Art. 109-116 CFPF), \
guarda y custodia, filiación, violencia familiar
- Derecho Civil: contratos, obligaciones, derechos reales, sucesiones, usucapión
- Derecho Penal: tipos penales, medidas cautelares, etapas procesales
- Derecho Laboral: despido, beneficios sociales, AFP, reincorporación
- Derecho Comercial: sociedades, títulos valores, quiebra
- Derecho Administrativo y Constitucional
- Derecho Tributario

ESTILO DE RESPUESTA
- Sé claro, preciso y profesional pero accesible
- Cita siempre el artículo y la ley específica cuando des información legal
- Si la consulta requiere análisis del caso concreto, pide los datos relevantes
- Advierte cuando una situación requiere atención presencial urgente
- Nunca reemplaces el criterio del abogado — eres un asistente de apoyo
- Responde siempre en español
- Usa formato limpio: párrafos cortos, listas cuando corresponda
- Si no tienes certeza sobre algo, indícalo claramente

CONTEXTO DEL DESPACHO
Trabajas para un despacho jurídico boliviano digital e innovador. \
Los usuarios son abogados, asistentes legales y clientes del despacho.
"""

@login_required
def api_chatbot(request):
    """
    Recibe: POST con JSON { "historial": [...], "mensaje": "..." }
    historial: lista de {"role": "user"|"model", "parts": "texto"}
    Devuelve: JSON { "respuesta": "...", "ok": true }
    """
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    try:
        body    = json.loads(request.body)
        mensaje = body.get('mensaje', '').strip()
        historial = body.get('historial', [])

        if not mensaje:
            return JsonResponse({'ok': False, 'error': 'Mensaje vacío'}, status=400)

        genai.configure(api_key=settings.GEMINI_API_KEY)
        modelo = genai.GenerativeModel(
            model_name='gemini-2.5-flash',
            system_instruction=SYSTEM_PROMPT_CHATBOT,
        )

        # Construir historial de conversación para Gemini
        chat = modelo.start_chat(history=historial)
        respuesta = chat.send_message(mensaje)

        return JsonResponse({
            'ok':       True,
            'respuesta': respuesta.text,
        })

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=500)
    
# ══════════════════════════════════════════════════════════════
# ASISTENTE IA POR EXPEDIENTE — Gemini 2.5 Flash
# A diferencia del LexBot general, este endpoint recibe el
# contexto completo del expediente (partes, movimientos,
# documentos, tareas) y lo inyecta en el system prompt.
# El abogado puede pedir redacciones, análisis y estrategias
# con datos reales del caso sin escribir nada extra.
# ══════════════════════════════════════════════════════════════

SYSTEM_PROMPT_ASISTENTE_EXPEDIENTE = """
Eres el Asistente Jurídico IA del despacho LexNova, especializado en \
Derecho boliviano. Tienes acceso al expediente completo que se detalla \
a continuación y debes usarlo como base de todas tus respuestas.

CAPACIDADES PRINCIPALES
1. ANÁLISIS DEL CASO: identifica fortalezas, debilidades y riesgos procesales
2. REDACCIÓN DE ESCRITOS: memoriales, solicitudes, conminatorias y recursos
   usando los datos reales del expediente (partes, juzgado, NUREJ)
3. PLAZOS PROCESALES: calcula y advierte sobre vencimientos según la norma aplicable
4. ESTRATEGIA LEGAL: sugiere próximos pasos según el estado del proceso
5. CITAS NORMATIVAS: referencia artículos exactos de la ley boliviana aplicable

LEGISLACIÓN QUE DOMINAS
- Constitución Política del Estado Plurinacional
- Código Procesal Civil (Ley 439) — plazos, recursos, audiencias
- Ley N° 603 CFPF — familia, asistencia familiar, divorcio, guarda
- Código Penal y Código de Procedimiento Penal (Ley 1970)
- Ley General del Trabajo y DS 224
- Ley 348 — violencia contra la mujer
- Código Civil (D.L. 12760) — contratos, obligaciones, derechos reales
- Ley 439 — nulidades, excepciones, recursos ordinarios y extraordinarios

REGLAS DE REDACCIÓN DE ESCRITOS
- Encabezado: SEÑOR JUEZ [juzgado del expediente]
- Identificación: nombre del cliente, número de expediente y NUREJ reales
- Cuerpo: hechos, fundamentos de derecho con artículos exactos, petitorio
- Cierre: "Es justicia que espero" / "Así lo pido de conformidad a Ley"
- Usa [DATO FALTANTE] para campos que no están en el expediente

ESTILO
- Profesional, preciso, directo
- Párrafos cortos, negritas para artículos legales clave
- Si detectás un plazo vencido o próximo a vencer, avisa de inmediato
- Si la consulta es ambigua, pregunta antes de redactar

CONTEXTO DEL EXPEDIENTE ACTUAL:
{contexto}
"""

@login_required
def api_asistente_expediente(request, pk):
    """
    Asistente IA por expediente con lectura real de documentos adjuntos.

    ESTRATEGIA DE LECTURA POR TIPO DE ARCHIVO:
    1. PDF con texto  → pypdf extrae texto directamente
    2. PDF escaneado  → se convierte a imágenes y Gemini las lee con visión
    3. .docx          → python-docx extrae texto
    4. Imágenes       → se envían directamente a Gemini como parte multimodal

    POST { "mensaje": "...", "historial": [...], "leer_docs": true/false }
    → JSON { "ok": true, "respuesta": "...", "docs_leidos": [...] }
    """
    if request.method != 'POST':
        return JsonResponse({'ok': False, 'error': 'Método no permitido'}, status=405)

    expediente = get_object_or_404(Expediente, pk=pk)

    try:
        body      = json.loads(request.body)
        mensaje   = body.get('mensaje', '').strip()
        historial = body.get('historial', [])
        leer_docs = body.get('leer_docs', False)  # el front lo activa si el usuario lo pide

        if not mensaje:
            return JsonResponse({'ok': False, 'error': 'Mensaje vacío'}, status=400)

        # ── 1. Construir contexto del expediente ──────────────
        from .models import Expediente as Exp
        tareas = expediente.tareas.select_related('asignada_a').order_by('fecha_limite')

        movimientos_texto = '\n'.join([
            f"- {m.fecha.strftime('%d/%m/%Y')} | {m.get_tipo_display()}: {m.descripcion}"
            for m in expediente.movimientos.all()[:10]
        ]) or 'Sin movimientos registrados.'

        documentos_lista = expediente.documentos.all()[:10]
        documentos_texto = '\n'.join([
            f"- {d.titulo} ({d.get_tipo_display()}) — {d.fecha_subida.strftime('%d/%m/%Y')}"
            for d in documentos_lista
        ]) or 'Sin documentos adjuntos.'

        tareas_texto = '\n'.join([
            f"- {'⚠️ VENCIDA' if t.esta_vencida else 'Vence'} "
            f"{t.fecha_limite.strftime('%d/%m/%Y')}: {t.titulo} [{t.get_estado_display()}]"
            for t in tareas
        ]) or 'Sin tareas asignadas.'

        contexto_base = f"""N° Expediente : {expediente.numero_expediente or '—'}
NUREJ / IANUS : {expediente.nurej or '—'}
Juzgado       : {expediente.juzgado or '—'}
Materia       : {expediente.materia}
Tipo de proceso: {expediente.tipo_proceso or '—'}
Estado        : {expediente.get_estado_display()}
Fecha inicio  : {expediente.fecha_inicio.strftime('%d/%m/%Y')}
Cliente       : {expediente.cliente.nombre_completo}
Abogado       : {str(expediente.abogado_asignado) if expediente.abogado_asignado else '—'}
Descripción   : {expediente.descripcion or '—'}

MOVIMIENTOS PROCESALES:
{movimientos_texto}

DOCUMENTOS ADJUNTOS:
{documentos_texto}

TAREAS:
{tareas_texto}"""

        system = SYSTEM_PROMPT_ASISTENTE_EXPEDIENTE.format(contexto=contexto_base)

        genai.configure(api_key=settings.GEMINI_API_KEY)

        # ── 2. Leer contenido real de los documentos ──────────
        # Se activa cuando el usuario pregunta sobre documentos
        # o cuando leer_docs=True viene en el request
        docs_leidos    = []
        partes_mensaje = []  # contenido multimodal para Gemini

        palabras_clave_docs = [
            'documento', 'adjunto', 'pdf', 'expediente', 'archivo',
            'resolución', 'auto', 'demanda', 'memorial', 'lee', 'leer',
            'analiza', 'analizar', 'revisar', 'revisa', 'dice', 'contiene',
            'qué dice', 'qué hay', 'muéstrame', 'mostrar',
        ]
        menciona_docs = leer_docs or any(
            p in mensaje.lower() for p in palabras_clave_docs
        )

        if menciona_docs and documentos_lista:
            texto_extraido_total = []

            for doc in documentos_lista[:5]:  # máximo 5 documentos por request
                try:
                    archivo_path = doc.archivo.path
                    ext = archivo_path.rsplit('.', 1)[-1].lower() if '.' in archivo_path else ''
                    nombre_doc = f"{doc.titulo} ({doc.get_tipo_display()})"

                    # ── PDF ──────────────────────────────────
                    if ext == 'pdf':
                        texto = _extraer_texto_pdf(archivo_path)

                        if texto and len(texto.strip()) > 100:
                            # PDF con texto digital — funciona perfecto
                            texto_extraido_total.append(
                                f"\n{'='*50}\n"
                                f"DOCUMENTO: {nombre_doc}\n"
                                f"{'='*50}\n"
                                f"{texto[:8000]}"
                            )
                            docs_leidos.append({'nombre': nombre_doc, 'metodo': 'texto'})
                        else:
                            # PDF escaneado — enviar páginas como imágenes a Gemini
                            imagenes_b64 = _pdf_a_imagenes_base64(archivo_path, max_paginas=6)
                            if imagenes_b64:
                                partes_mensaje.append(
                                    f"\n[DOCUMENTO ESCANEADO: {nombre_doc}]"
                                    f"\nAnálisis de {len(imagenes_b64)} página(s):"
                                )
                                for img_data in imagenes_b64:
                                    partes_mensaje.append({
                                        'inline_data': {
                                            'mime_type': 'image/jpeg',
                                            'data':      img_data,
                                        }
                                    })
                                docs_leidos.append({'nombre': nombre_doc, 'metodo': 'vision'})
                            else:
                                texto_extraido_total.append(
                                    f"\n[{nombre_doc}: PDF sin texto legible ni imágenes extraíbles]"
                                )
                                docs_leidos.append({'nombre': nombre_doc, 'metodo': 'no_leido'})

                    # ── DOCX ─────────────────────────────────
                    elif ext == 'docx':
                        texto = _extraer_texto_docx(archivo_path)
                        if texto:
                            texto_extraido_total.append(
                                f"\n{'='*50}\n"
                                f"DOCUMENTO: {nombre_doc}\n"
                                f"{'='*50}\n"
                                f"{texto[:8000]}"
                            )
                            docs_leidos.append({'nombre': nombre_doc, 'metodo': 'texto'})

                    # ── IMÁGENES ──────────────────────────────
                    elif ext in ('jpg', 'jpeg', 'png', 'webp'):
                        import base64
                        with open(archivo_path, 'rb') as f:
                            img_b64 = base64.b64encode(f.read()).decode('utf-8')
                        mime = 'image/jpeg' if ext in ('jpg', 'jpeg') else f'image/{ext}'
                        partes_mensaje.append(f"\n[IMAGEN: {nombre_doc}]")
                        partes_mensaje.append({
                            'inline_data': {'mime_type': mime, 'data': img_b64}
                        })
                        docs_leidos.append({'nombre': nombre_doc, 'metodo': 'vision'})

                except Exception as e:
                    docs_leidos.append({'nombre': doc.titulo, 'metodo': 'error', 'detalle': str(e)})
                    continue

            # Agregar texto extraído al mensaje
            if texto_extraido_total:
                partes_mensaje.insert(0,
                    f"CONTENIDO DE LOS DOCUMENTOS ADJUNTOS:\n"
                    f"{''.join(texto_extraido_total)}\n\n"
                    f"CONSULTA DEL ABOGADO: {mensaje}"
                )
            else:
                partes_mensaje.insert(0, mensaje)
        else:
            partes_mensaje = [mensaje]

        # ── 3. Llamar a Gemini ────────────────────────────────
        modelo = genai.GenerativeModel(
            model_name='gemini-2.5-flash',
            system_instruction=system,
        )
        chat      = modelo.start_chat(history=historial)
        respuesta = chat.send_message(partes_mensaje)

        # Actualizar historial (solo texto, no imágenes)
        historial.append({'role': 'user',  'parts': mensaje})
        historial.append({'role': 'model', 'parts': respuesta.text})
        if len(historial) > 16:
            historial = historial[-16:]

        return JsonResponse({
            'ok':         True,
            'respuesta':  respuesta.text,
            'docs_leidos': docs_leidos,
            'historial':   historial,
        })

    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=500)


# ── Helpers de extracción de archivos ────────────────────────

def _extraer_texto_pdf(ruta: str) -> str:
    """Extrae texto de PDF digitales con pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(ruta)
        return '\n'.join(
            (page.extract_text() or '') for page in reader.pages[:25]
        )
    except Exception:
        return ''


def _extraer_texto_docx(ruta: str) -> str:
    """Extrae texto de archivos .docx."""
    try:
        from docx import Document
        doc = Document(ruta)
        return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        return ''


def _pdf_a_imagenes_base64(ruta: str, max_paginas: int = 6) -> list[str]:
    """
    Convierte páginas de un PDF escaneado a imágenes base64.
    Requiere pdf2image + poppler. Si no están instalados, retorna [].
    """
    try:
        import base64
        from pdf2image import convert_from_path
        paginas = convert_from_path(ruta, dpi=150, last_page=max_paginas)
        resultado = []
        for pagina in paginas:
            buf = io.BytesIO()
            pagina.save(buf, format='JPEG', quality=75)
            resultado.append(base64.b64encode(buf.getvalue()).decode('utf-8'))
        return resultado
    except Exception:
        return []

# ── IMPORTACIONES ADICIONALES NECESARIAS EN views.py ──────────
@login_required
@require_POST
def api_explicar_idioma_nativo(request):
    """
    Endpoint AJAX — recibe texto + idioma + contexto,
    devuelve explicación en Quechua o Aymara + español simple.

    POST params:
        texto    : texto jurídico a explicar
        idioma   : 'quechua' | 'aymara'
        contexto : 'expediente' | 'movimiento' | 'documento' |
                   'contrato' | 'memorial' | 'liquidacion' |
                   'tarea' | 'evento' | 'jurisprudencia' | 'general'
    """
    texto    = request.POST.get('texto', '').strip()
    idioma   = request.POST.get('idioma', 'quechua').strip().lower()
    contexto = request.POST.get('contexto', 'general').strip().lower()

    if not texto:
        return JsonResponse({'error': 'No se recibió texto para explicar.'}, status=400)
    if idioma not in ('quechua', 'aymara'):
        return JsonResponse({'error': 'Idioma no válido. Usa quechua o aymara.'}, status=400)

    explicacion, error = explicar_en_idioma_nativo(
        texto=texto,
        idioma=idioma,
        contexto=contexto,
    )

    if error:
        return JsonResponse({'error': error}, status=500)

    return JsonResponse({
        'explicacion': explicacion,
        'idioma':      idioma,
        'contexto':    contexto,
    })

@login_required
def calculadora_beneficios_sociales(request):
    """
    Calculadora de Beneficios Sociales y Derechos Laborales — Derecho Laboral Boliviano.
    LGT · D.S. 110 · D.S. 28699 · D.S. 522 · D.S. 21060 · Ley 18/12/1944
    El cálculo es 100% en el frontend (JS). La vista solo renderiza el template.
    """
    return render(request, 'control_legal/beneficios_sociales.html') 

# ══════════════════════════════════════════════════════════════
# REPORTE PDF — PREFINIQUITO LABORAL
# Recibe los datos calculados en JS via POST y genera el PDF
# ══════════════════════════════════════════════════════════════
@login_required
def exportar_prefiniquito_pdf(request):
    if request.method != 'POST':
        return redirect('calculadora_beneficios_sociales')
    try:
        from .prefiniquito_pdf_services import generar_pdf_prefiniquito
        datos = json.loads(request.body)
        buf   = generar_pdf_prefiniquito(datos)
        nombre_trabajador = datos.get('trabajador', 'prefiniquito').replace(' ', '_')
        response = HttpResponse(buf.read(), content_type='application/pdf')
        response['Content-Disposition'] = (
            f'inline; filename="Prefiniquito_{nombre_trabajador}.pdf"'
        )
        return response
    except Exception as e:
        return HttpResponse(f'Error generando PDF: {str(e)}', status=500)